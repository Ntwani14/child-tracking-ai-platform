from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins (tighter to save space) ──────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# ── Helpers ───────────────────────────────────────────────────
def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(doc, text):
    add_divider(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    set_font(run, size=10.5, bold=True, color=(31, 56, 100))

def job_title_line(doc, role, company, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(role)
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {dates}")
    set_font(r2, size=9.5, color=(89, 89, 89))

def bullet(doc, text, indent=0.2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(text)
    set_font(run, size=9.5)

def normal(doc, text, size=9.5, bold=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)

# ══════════════════════════════════════════════════════════════
# NAME & CONTACT
# ══════════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(1)
nr = name_p.add_run("NTWANANO MATHEBULA")
set_font(nr, size=18, bold=True, color=(31, 56, 100))

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
tr = title_p.add_run("QA Automation Engineer  |  Technical Test Analyst")
set_font(tr, size=11, color=(89, 89, 89))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
cr = contact_p.add_run(
    "Johannesburg, Gauteng, SA   |   064 064 5892   |   "
    "ntwano14@gmail.com   |   linkedin.com/in/ntwanano-mathebula-7bb152196"
)
set_font(cr, size=9.5, color=(89, 89, 89))

# ══════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Professional Summary")
normal(doc,
    "QA Automation Engineer and Technical Test Analyst with over 3 years of experience in "
    "software quality assurance and test automation within the banking and financial technology "
    "sectors. Currently employed at FNB South Africa, specialising in automation of Card Banking "
    "applications including mobile, web, SOAP, and REST API testing. Experienced in building and "
    "maintaining UI and API automation frameworks using Java, Playwright, Selenium WebDriver, and "
    "REST Assured. Skilled in performance testing, database validation, defect management, and "
    "Agile Scrum delivery. Strong understanding of OOP principles with proven ability to automate "
    "business-critical financial systems.",
    size=9.5
)

# ══════════════════════════════════════════════════════════════
# CORE COMPETENCIES
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Core Competencies")
competencies = (
    "UI & API Test Automation  |  Mobile & Web Application Testing  |  Functional & Regression Testing  |  "
    "Integration & End-to-End Testing  |  Performance & Load Testing  |  SOAP and RESTful Services  |  "
    "OOP Principles  |  Test Planning & Case Design  |  Defect Management  |  Requirements Analysis  |  "
    "Banking Systems Testing  |  Agile Scrum / CI/CD  |  Building Partnerships  |  High-Impact Communication"
)
normal(doc, competencies, size=9.5, space_after=2)

# ══════════════════════════════════════════════════════════════
# TECHNICAL SKILLS
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Technical Skills")

skills = [
    ("Programming Languages", "Java, SQL, JavaScript, HTML, CSS"),
    ("Automation Frameworks", "Playwright, Selenium WebDriver, Cucumber BDD, REST Assured, Appium (Mobile), JUnit, TestNG"),
    ("API Testing",           "REST APIs, SOAP Web Services, SoapUI, Postman, Karate Framework"),
    ("Performance Testing",   "Apache JMeter"),
    ("Databases",             "PostgreSQL, MySQL, Microsoft SQL Server"),
    ("DevOps & Version Control", "Jenkins, Bitbucket, Git, Maven"),
    ("Test Management",       "Jira"),
    ("Development",           "Spring Boot, Hibernate, Swagger/OpenAPI"),
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r_label = p.add_run(f"{label}: ")
    set_font(r_label, size=9.5, bold=True)
    r_value = p.add_run(value)
    set_font(r_value, size=9.5)

# ══════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Professional Experience")

# FNB
job_title_line(doc, "Technical Test Analyst", "FNB South Africa", "Johannesburg", "March 2025 – Present")
fnb_bullets = [
    "Perform functional, system, integration, regression, and end-to-end testing for Card Banking applications across Mobile App, AOP, and OBE platforms.",
    "Assigned to the Fleet project, validating business-critical functionality across mobile, web, SOAP services, and REST APIs.",
    "Develop and maintain UI automation scripts using Java with Playwright (migrated from Selenium WebDriver).",
    "Automate SOAP and RESTful API test suites using the bank's internal framework; validate using SoapUI and Postman.",
    "Write SQL scripts to validate backend data and confirm consistency between application and database layers.",
    "Analyse business requirements to design comprehensive test cases aligned to acceptance criteria.",
    "Log and manage defects in Jira; collaborate with developers, BAs, and product owners to drive resolution.",
    "Execute automated pipelines via Jenkins; manage source code using Bitbucket.",
    "Participate in all Agile Scrum ceremonies: sprint planning, backlog refinement, stand-ups, and retrospectives.",
]
for b in fnb_bullets:
    bullet(doc, b)

# CSIR
job_title_line(doc, "Software Developer Intern", "CSIR", "Pretoria", "September 2024 – March 2025")
for b in [
    "Developed RESTful APIs using Java, Spring Boot, and PostgreSQL following REST architecture principles.",
    "Documented APIs using Swagger/OpenAPI; tested with Postman.",
    "Wrote unit tests using JUnit and Mockito.",
    "Participated in code reviews, debugging, and solution design within an Agile Scrum team.",
]:
    bullet(doc, b)

# Deviare
job_title_line(doc, "Full Stack Java Developer", "Deviare", "Johannesburg", "September 2023 – August 2024")
for b in [
    "Developed backend applications using Java, Spring Boot, Hibernate, and PostgreSQL.",
    "Designed responsive frontend interfaces using HTML, CSS, JavaScript, and Bootstrap.",
    "Developed and integrated RESTful APIs; wrote unit tests using JUnit and Mockito.",
    "Collaborated within Agile development teams throughout the full SDLC.",
]:
    bullet(doc, b)

# Nagarro
job_title_line(doc, "QA Automation Engineer", "Nagarro", "Johannesburg", "April 2022 – April 2023")
for b in [
    "Designed and maintained UI automation frameworks using Selenium WebDriver with Java.",
    "Automated REST API testing using REST Assured; implemented BDD automation using Cucumber.",
    "Developed reusable test scripts covering functional, regression, smoke, sanity, and integration testing.",
    "Conducted performance and load testing using Apache JMeter.",
    "Performed backend database validation using SQL; validated API responses using Postman.",
    "Logged and tracked defects in Jira; participated in Agile Scrum ceremonies.",
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# PROJECTS
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Projects")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Waste Sorting Backend Application")
set_font(r, size=9.5, bold=True)
normal(doc, "Technologies: Java, Spring Boot, REST APIs, H2 Database", size=9.5, space_after=1)
for b in [
    "Developed RESTful APIs for waste category management with full CRUD operations.",
    "Applied layered architecture and OOP design principles; tested endpoints using Postman.",
]:
    bullet(doc, b)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(3)
p2.paragraph_format.space_after  = Pt(1)
r2 = p2.add_run("Vehicle Tracking System")
set_font(r2, size=9.5, bold=True)
normal(doc, "Technologies: Java, Spring Boot, Hibernate, PostgreSQL, HTML, CSS, JavaScript, Bootstrap", size=9.5, space_after=1)
for b in [
    "Developed backend services for vehicle registration and tracking with REST API integration.",
    "Built responsive user interface using Bootstrap; performed unit testing using JUnit.",
]:
    bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# EDUCATION
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Education")
edu = [
    ("Advanced Diploma in Computer Science", "Tshwane University of Technology", "2025 – 2026"),
    ("National Diploma in Information Technology", "Tshwane University of Technology", "2017 – 2020"),
]
for qual, inst, period in edu:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(qual)
    set_font(r1, size=9.5, bold=True)
    r2 = p.add_run(f"  —  {inst}  |  {period}")
    set_font(r2, size=9.5)

# ══════════════════════════════════════════════════════════════
# CERTIFICATIONS
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Certifications")
certs = [
    "SQL for Testers",
    "Java Object-Oriented Programming",
    "Microsoft Security, Compliance, and Identity Fundamentals (SC-900)",
    "Virtual Global Apprenticeship in Software Development and Engineering",
]
for c in certs:
    bullet(doc, c)

# ══════════════════════════════════════════════════════════════
# LANGUAGES
# ══════════════════════════════════════════════════════════════
section_heading(doc, "Languages")
normal(doc, "English  |  Xitsonga", size=10.5)

# ── Save ──────────────────────────────────────────────────────
output_path = r"c:\child-tracking-ai-platform\Ntwanano_Mathebula_CV.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
